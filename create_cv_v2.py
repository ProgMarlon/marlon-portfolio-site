from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(doc, top=0.75, bottom=0.75, left=0.75, right=0.75):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Professional Blue
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    # Add a border bottom
    p_element = p._element
    pPr = p_element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pbdr.append(bottom)
    pPr.append(pbdr)

def create_cv():
    doc = Document()
    set_margins(doc)

    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # Name Header
    header_name = doc.add_paragraph()
    header_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_name.add_run('MARLON C. ISAGUIRRE JR.')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 41, 55)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Dasmariñas, Cavite, Philippines | +63 994 818 0273\n')
    contact.add_run('marloncopilot@gmail.com | ').font.color.rgb = RGBColor(37, 99, 235)
    contact.add_run('github.com/ProgMarlon | ').font.color.rgb = RGBColor(37, 99, 235)
    contact.add_run('linkedin.com/in/marlon-isaguirre-jr-414683393').font.color.rgb = RGBColor(37, 99, 235)

    # Professional Summary
    add_section_header(doc, 'Professional Summary')
    summary = doc.add_paragraph(
        "Results-oriented Full-Stack Web Developer and 3rd-year BSIT student with expertise in the MERN stack and Python. "
        "Proven ability to engineer high-performance applications, recently achieving a 20% reduction in load times through "
        "strategic optimization. Passionate about bridging technical complexity with intuitive user experiences and "
        "maintaining high accessibility standards (WCAG AA)."
    )
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Technical Skills
    add_section_header(doc, 'Technical Skills')
    skills_table = doc.add_table(rows=0, cols=2)
    skills_table.autofit = True
    
    skills = [
        ("Frontend", "React.js, TypeScript, Next.js, HTML5, CSS3, Tailwind, Bootstrap"),
        ("Backend", "Node.js, Express, Python (FastAPI/Flask), PHP, RESTful APIs"),
        ("Databases", "MongoDB, MySQL, PostgreSQL"),
        ("Tools/DevOps", "Git, GitHub, Vercel, VS Code, CI/CD Pipelines, Docker"),
        ("Practices", "Agile/Scrum, Unit Testing (Jest), SEO, WCAG AA Accessibility")
    ]
    
    for category, items in skills:
        row_cells = skills_table.add_row().cells
        row_cells[0].text = f"• {category}:"
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = items

    # Professional Experience
    add_section_header(doc, 'Professional Experience')
    
    # Freelance
    exp1 = doc.add_paragraph()
    run = exp1.add_run('FULL-STACK DEVELOPER (Freelance)')
    run.bold = True
    run.font.size = Pt(12)
    exp1.add_run('\nIndependent Projects | Jan 2023 – Present').italic = True
    
    doc.add_paragraph("Engineered custom MERN stack solutions for diverse clients, focusing on scalability and user retention.", style='List Bullet')
    doc.add_paragraph("Achieved 20% improvement in page performance metrics via image optimization and lazy loading.", style='List Bullet')
    doc.add_paragraph("Integrated complex 3rd-party APIs (Google Maps, Mailers) to enhance business automation.", style='List Bullet')

    # Academic Leadership
    exp2 = doc.add_paragraph()
    run = exp2.add_run('TECHNICAL LEAD & ARCHITECT')
    run.bold = True
    run.font.size = Pt(12)
    exp2.add_run('\nMajor Development Projects | Sep 2024 – Present').italic = True
    
    doc.add_paragraph("Spearheaded the architectural design of AI-driven systems and microservices platforms.", style='List Bullet')
    doc.add_paragraph("Implemented CI/CD workflows using GitHub Actions to ensure reliable code deployment.", style='List Bullet')
    doc.add_paragraph("Led cross-functional academic teams using Agile methodologies to deliver projects on strict deadlines.", style='List Bullet')

    # Selected Projects
    add_section_header(doc, 'Key Technical Projects')
    
    # MarBot
    p1 = doc.add_paragraph()
    p1.add_run('MarBot2000 (AI Virtual Assistant)').bold = True
    p1.add_run(' | React, Node.js, Gemini AI').italic = True
    doc.add_paragraph("Built an intelligent assistant using Gemini 2.5 Flash with persistent LocalStorage history.", style='List Bullet')
    doc.add_paragraph("Designed an accessible, retro-themed interface with WCAG AA compliance and 100% responsiveness.", style='List Bullet')

    # StockSense
    p2 = doc.add_paragraph()
    p2.add_run('StockSense (Predictive Inventory System)').bold = True
    p2.add_run(' | React, FastAPI, MongoDB').italic = True
    doc.add_paragraph("Developed a microservices dashboard that predicts stock needs based on historical sales trends.", style='List Bullet')
    doc.add_paragraph("Created dynamic data visualizations using Recharts, allowing for instant data-driven stock analysis.", style='List Bullet')

    # Hackathons
    p3 = doc.add_paragraph()
    p3.add_run('PowerHabit (Meralco IDOL Hackathon 2024)').bold = True
    doc.add_paragraph("Rapidly prototyped a cost-tracking app for non-smart appliance users to monitor energy savings.", style='List Bullet')

    # Education
    add_section_header(doc, 'Education')
    edu = doc.add_paragraph()
    edu.add_run('Bachelor of Science in Information Technology').bold = True
    edu.add_run('\nCavite State University — Don Severino de las Alas Campus | 2023 – Present')
    doc.add_paragraph("Dean's List / Consistent academic high-performer. Focus: Software Engineering.", style='List Bullet')

    # Certifications & Awards
    add_section_header(doc, 'Certifications & Awards')
    certs = [
        "Responsive Web Design Certification | freeCodeCamp | 2024",
        "JavaScript Algorithms and Data Structures | freeCodeCamp | 2023",
        "Git and GitHub Essentials | Great Learning Academy | 2024",
        "Participant, DLSU Hackercup 2025"
    ]
    for cert in certs:
        doc.add_paragraph(cert, style='List Bullet')

    # Languages & References
    add_section_header(doc, 'Additional Information')
    lang = doc.add_paragraph()
    lang.add_run('Languages: ').bold = True
    lang.add_run('English (Professional), Tagalog (Native)')
    
    ref = doc.add_paragraph()
    ref.add_run('References: ').bold = True
    ref.add_run('Available upon request.')

    # Save
    import os
    output_path = os.path.join(os.getcwd(), 'Marlon_Isaguirre_Professional_CV.docx')
    doc.save(output_path)
    print(f"Professional CV created at: {output_path}")

if __name__ == "__main__":
    create_cv()
