from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_cv():
    doc = Document()

    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Header
    header = doc.add_heading('Marlon C. Isaguirre Jr.', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Dasmariñas, Cavite, Philippines | +63 994 818 0273')
    contact.add_run('\nmarloncopilot@gmail.com | github.com/ProgMarlon | x.com/MarBot2000')

    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        "Results-driven Web Developer and 3rd-year BSIT student with a strong focus on building high-performance, "
        "accessible, and user-centric applications. Specializing in the MERN stack and Python, I excel at bridging "
        "technical complexity with seamless user experiences. Proven track record in optimizing deployment workflows "
        "and integrating AI solutions to deliver measurable business value."
    )

    # Experience
    doc.add_heading('Experience', level=1)
    
    exp1 = doc.add_paragraph()
    exp1.add_run('Full-Stack Developer (Freelance)').bold = True
    exp1.add_run('\nIndependent Projects | 2023 – Present')
    doc.add_paragraph(
        "• Delivering custom MERN stack solutions for diverse clients with a focus on performance.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Achieving up to 20% reduction in load times through image optimization and lazy loading.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Integrating third-party APIs and building secure backend infrastructures.",
        style='List Bullet'
    )

    exp2 = doc.add_paragraph()
    exp2.add_run('Technical Lead & Architect').bold = True
    exp2.add_run('\nAcademic Projects | 2024 – 2025')
    doc.add_paragraph(
        "• Led development of complex systems like StockSense and MarBot2000.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Managed the transition from monolithic to microservices architecture for improved scalability.",
        style='List Bullet'
    )
    doc.add_paragraph(
        "• Orchestrated collaborative workflows using Git and Agile methodologies.",
        style='List Bullet'
    )

    # Projects
    doc.add_heading('Key Projects', level=1)
    
    doc.add_paragraph('MarBot2000 (AI Virtual Assistant)', style='List Bullet').bold = True
    doc.add_paragraph("• Built an intelligent chatbot integrated with Gemini 2.5 Flash for natural language interactions.", style='Normal')
    doc.add_paragraph("• Implemented persistent message history using LocalStorage for seamless session retrieval.", style='Normal')
    
    doc.add_paragraph('StockSense (Inventory Forecast System)', style='List Bullet').bold = True
    doc.add_paragraph("• Engineered a full-stack dashboard utilizing a Node.js API and a Python (FastAPI) analytics engine.", style='Normal')
    doc.add_paragraph("• Developed a predictive model to forecast inventory demand based on historical sales data.", style='Normal')

    doc.add_paragraph('PowerHabit (Meralco IDOL Hackathon 2024)', style='List Bullet').bold = True
    doc.add_paragraph("• Developed an application to help non-smart appliance users track energy costs and savings.", style='Normal')

    doc.add_paragraph('Parkada (DLSU Hackercup 2025)', style='List Bullet').bold = True
    doc.add_paragraph("• Built a parking reservation platform enabling users to book spaces via a user-friendly interface.", style='Normal')

    # Education
    doc.add_heading('Education', level=1)
    edu = doc.add_paragraph()
    edu.add_run('Bachelor of Science in Information Technology').bold = True
    edu.add_run('\nCavite State University — Don Severino de las Alas Campus | 2023 – Present')
    doc.add_paragraph("Focus: Web Development and Software Engineering.", style='List Bullet')

    # Skills
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph(
        "• Languages: JavaScript (ES6+), TypeScript, Python, Java, PHP, SQL\n"
        "• Technologies: React.js, Node.js, Express, MongoDB, FastAPI, Bootstrap, Tailwind CSS\n"
        "• Tools: Git, GitHub, Vercel, VS Code, REST APIs, Google Maps API\n"
        "• Practices: CI/CD, Unit Testing, Agile, Performance Optimization, WCAG AA Accessibility",
        style='Normal'
    )

    # Certifications
    doc.add_heading('Certifications', level=1)
    doc.add_paragraph(
        "• Responsive Web Design Certification | freeCodeCamp | 2024\n"
        "• JavaScript Algorithms and Data Structures | freeCodeCamp | 2023\n"
        "• Git and GitHub Essentials | Great Learning Academy | 2024",
        style='Normal'
    )

    # Awards & Achievements
    doc.add_heading('Awards & Achievements', level=1)
    doc.add_paragraph(
        "• Participant, Meralco IDOL Hackathon 2024\n"
        "• Participant, DLSU Hackercup 2025\n"
        "• Successfully optimized portfolio performance to 90+ Lighthouse scores.",
        style='Normal'
    )

    # Languages
    doc.add_heading('Languages', level=1)
    doc.add_paragraph("English (Professional Proficiency), Tagalog (Native)", style='Normal')

    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph("Available upon request.", style='Normal')

    # Save
    import os
    output_path = os.path.join(os.getcwd(), 'Marlon_Isaguirre_CV.docx')
    doc.save(output_path)
    print(f"CV created successfully at: {output_path}")

if __name__ == "__main__":
    create_cv()